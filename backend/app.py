import os
import tempfile
import logging
import uuid
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from dotenv import load_dotenv

from PyPDF2 import PdfReader, PdfWriter
import docx
import google.generativeai as genai
from google.cloud import documentai_v1 as documentai

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
PROJECT_ID = os.getenv("GCP_PROJECT_ID")
PROCESSOR_ID = os.getenv("DOC_AI_PROCESSOR_ID")
LOCATION = os.getenv("DOC_AI_LOCATION", "us")

if not GOOGLE_API_KEY:
    raise RuntimeError("Bạn phải thiết lập biến môi trường GOOGLE_API_KEY")

genai.configure(api_key=GOOGLE_API_KEY)

app = FastAPI(title="Doc Summarizer (Gemini + Document AI, Export .docx/.txt)")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
logger = logging.getLogger("uvicorn.error")

def split_pdf(input_pdf, pages_per_chunk=15):
    reader = PdfReader(input_pdf)
    total_pages = len(reader.pages)
    chunk_files = []
    for i in range(0, total_pages, pages_per_chunk):
        writer = PdfWriter()
        for j in range(i, min(i + pages_per_chunk, total_pages)):
            writer.add_page(reader.pages[j])
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            writer.write(temp_file)
            chunk_files.append(temp_file.name)
    return chunk_files

def extract_text_docai(file_path, mime_type="application/pdf"):
    client = documentai.DocumentProcessorServiceClient()
    name = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"
    with open(file_path, "rb") as f:
        content = f.read()
    raw_document = documentai.RawDocument(content=content, mime_type=mime_type)
    request = documentai.ProcessRequest(
        name=name,
        raw_document=raw_document
    )
    result = client.process_document(request=request)
    return result.document.text

def extract_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def summarize_text_gemini(text, language="en"):
    language_map = {
        "en": "English",
        "vi": "Vietnamese",
        "fr": "French",
        "zh": "Chinese",
        "ja": "Japanese"
    }
    lang_prompt = language_map.get(language, "English")
    prompt = (
        f"Summarize the following document in {lang_prompt}:\n\n"
        f"{text}"
    )
    max_input = 48000
    prompt = prompt[:max_input]
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    return response.text

def save_summary_to_docx(summary_text):
    doc = docx.Document()
    doc.add_paragraph(summary_text)
    filename = f"summary_{uuid.uuid4().hex}.docx"
    filepath = f"/tmp/{filename}"
    doc.save(filepath)
    return filename, filepath

def save_summary_to_txt(summary_text):
    filename = f"summary_{uuid.uuid4().hex}.txt"
    filepath = f"/tmp/{filename}"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(summary_text)
    return filename, filepath

@app.get("/")
async def root():
    return {"msg": "backend running"}

@app.post("/summarize/")
async def summarize(
    file: UploadFile = File(...),
    language: str = Form("en")
):
    allowed = {
        "pdf": "application/pdf",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in allowed:
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ PDF, DOCX, JPG, PNG")
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        if ext == "docx":
            text = extract_docx(tmp_path)
        elif ext == "pdf":
            chunk_paths = split_pdf(tmp_path, pages_per_chunk=15)
            all_text = []
            for idx, chunk in enumerate(chunk_paths):
                logger.info(f"OCR đoạn {idx+1}/{len(chunk_paths)}...")
                text_part = extract_text_docai(chunk)
                all_text.append(f"\n--- PHẦN {idx+1} ---\n{text_part}")
                os.remove(chunk)
            text = "\n".join(all_text)
        else:
            text = extract_text_docai(tmp_path, allowed[ext])
    except Exception as e:
        logger.error("Error extracting text: %s", e, exc_info=True)
        raise HTTPException(status_code=422, detail=f"Lỗi chi tiết: {e}")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Tài liệu rỗng hoặc OCR thất bại")

    try:
        summary = summarize_text_gemini(text, language=language)
    except Exception as e:
        logger.error("Error calling Gemini: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Lỗi khi gọi Gemini: {e}")

    return {"summary": summary}

@app.post("/export-summary/")
async def export_summary(
    summary: str = Form(...),
    export_format: str = Form("docx")  # docx hoặc txt
):
    if export_format.lower() == "txt":
        filename, filepath = save_summary_to_txt(summary)
        media_type = "text/plain"
    else:
        filename, filepath = save_summary_to_docx(summary)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(
        filepath,
        media_type=media_type,
        filename=filename
    )
